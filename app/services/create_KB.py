import os
import glob
import requests
from bs4 import BeautifulSoup
from docx import Document
import pandas as pd
import io
import re
import logging
import PyPDF2
import json
import tiktoken

# LangChain / FAISS
from langchain_community.vectorstores import FAISS
from langchain.schema import Document as LCDocument

logger = logging.getLogger(__name__)
logging.basicConfig(level=logging.INFO, format="%(message)s")

from app.core.config import FAISS_INDEX_PATH


class VectorStore:
    def __init__(self, embeddings, model_name: str = "sentence-transformers/all-MiniLM-L6-v2"):
        self.embeddings = embeddings
        self.vectors = None
        self.model_name = model_name
        # use tiktoken to tokenize safely
        self.encoding = tiktoken.get_encoding("cl100k_base")

    # -------- Tokenization --------
    def count_tokens(self, text: str) -> int:
        return len(self.encoding.encode(text))

    def split_by_tokens(self, text: str, max_tokens: int = 512, overlap: int = 50) -> list[str]:
        """Split text into chunks capped at max_tokens with overlap."""
        words = text.split()
        chunks, current_chunk = [], []

        for word in words:
            current_chunk.append(word)
            if self.count_tokens(" ".join(current_chunk)) >= max_tokens:
                chunks.append(" ".join(current_chunk))
                # keep overlap
                current_chunk = current_chunk[-overlap:]

        if current_chunk:
            chunks.append(" ".join(current_chunk))

        return chunks

    # -------- HTML --------
    def extract_html(self, url: str) -> str:
        response = requests.get(url)
        soup = BeautifulSoup(response.text, "html.parser")
        return soup.get_text(separator="\n", strip=True)

    # -------- PDF --------
    def extract_pdf(self, url: str) -> str:
        response = requests.get(url)
        with io.BytesIO(response.content) as f:
            reader = PyPDF2.PdfReader(f)
            text = "\n".join(
                page.extract_text() for page in reader.pages if page.extract_text()
            )
        return text

    # -------- DOCX --------
    def extract_docx(self, url: str) -> str:
        response = requests.get(url)
        with io.BytesIO(response.content) as f:
            doc = Document(f)
            text = "\n".join(p.text for p in doc.paragraphs)
        return text

    # -------- XLSX --------
    def extract_xlsx(self, url: str) -> str:
        df = pd.read_excel(url)
        return df.to_string()

    # -------- Google Doc (Public) --------
    def extract_google_doc(self, doc_url: str) -> str:
        file_id = re.findall(r"/d/([a-zA-Z0-9-_]+)", doc_url)[0]
        export_url = f"https://docs.google.com/document/d/{file_id}/export?format=txt"
        response = requests.get(export_url)
        return response.text

    # -------- Google Sheet (Public) --------
    def extract_google_sheet(self, sheet_url: str) -> str:
        sheet_id = re.findall(r"/d/([a-zA-Z0-9-_]+)", sheet_url)[0]
        export_url = f"https://docs.google.com/spreadsheets/d/{sheet_id}/export?format=csv"
        response = requests.get(export_url)
        df = pd.read_csv(io.StringIO(response.text))
        return df.to_string()

    # -------- Dispatcher --------
    def extract(self, url: str) -> str:
        if "docs.google.com/document" in url:
            return self.extract_google_doc(url)
        elif "docs.google.com/spreadsheets" in url:
            return self.extract_google_sheet(url)
        elif url.endswith(".pdf"):
            return self.extract_pdf(url)
        elif url.endswith(".docx"):
            return self.extract_docx(url)
        elif url.endswith(".xlsx"):
            return self.extract_xlsx(url)
        else:
            return self.extract_html(url)

    # -------- Vector Store Management --------
    def ensure_index(self, url: str) -> dict:
        # 🚨 Always start fresh (ignore old index if exists)
        if os.path.exists(FAISS_INDEX_PATH):
            import shutil
            shutil.rmtree(FAISS_INDEX_PATH)  # delete old vector store
            logger.info("🧹 Old FAISS index deleted. Creating new one.")

        # Extract text
        text = self.extract(url)
        logger.info("Text extracted successfully.")

        # Count total tokens
        total_tokens = self.count_tokens(text)

        # Split into chunks (512 max tokens)
        chunks = self.split_by_tokens(text, max_tokens=512, overlap=50)
        documents = [LCDocument(page_content=chunk) for chunk in chunks]

        # Count tokens after chunking
        embedded_tokens = sum(self.count_tokens(chunk) for chunk in chunks)

        # Create new FAISS
        self.vectors = FAISS.from_documents(documents, self.embeddings)
        os.makedirs(FAISS_INDEX_PATH, exist_ok=True)
        self.vectors.save_local(FAISS_INDEX_PATH)

        logger.info("✅ New vector embeddings created and saved.")
        logger.info(f"Total tokens in document: {total_tokens}")
        logger.info(f"Total tokens embedded: {embedded_tokens}")

        return self._make_metadata(url, "New vector embeddings created and saved.", total_tokens, embedded_tokens)


    def _make_metadata(self, url: str, status: str, total_tokens: int, embedded_tokens: int) -> dict:
        metadata = {
            "status": status,
            "source_url": url,
            "embedding_model": self.model_name,
            "vectorstore_path": FAISS_INDEX_PATH,
            "total_tokens": total_tokens,
            "embedded_tokens": embedded_tokens,
        }
        logger.info(json.dumps(metadata, indent=2))
        return metadata


# # --------- TEST ---------
# if __name__ == "__main__":
#     store = VectorStore()
#     test_url = "https://proceedings.neurips.cc/paper_files/paper/2017/file/3f5ee243547dee91fbd053c1c4a845aa-Paper.pdf"
#     result = store.ensure_index(test_url)
#     print(result)