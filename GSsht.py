# import gspread
# import logging
# from oauth2client.service_account import ServiceAccountCredentials
# from googleapiclient.discovery import build

# # -------------------------
# # Logger Setup
# # -------------------------
# logging.basicConfig(
#     level=logging.INFO,
#     format="%(asctime)s [%(levelname)s] %(message)s",
#     datefmt="%Y-%m-%d %H:%M:%S"
# )
# logger = logging.getLogger(__name__)

# def _connect_to_google_sheet(json_key_file, sheet_url):
#     try:
#         scope = [
#             "https://spreadsheets.google.com/feeds",
#             "https://www.googleapis.com/auth/drive"
#         ]
#         creds = ServiceAccountCredentials.from_json_keyfile_name(
#             json_key_file, scope
#         )
#         client = gspread.authorize(creds)
#         sheet = client.open_by_url(sheet_url).sheet1
#         logger.info("✅ Connected to Google Sheet successfully")
#         return sheet
#     except Exception as e:
#         logger.error(f"❌ Failed to connect to Google Sheet: {e}")
#         raise


# def _connect_to_google_doc(json_key_file, doc_id):
#     try:
#         scope = ["https://www.googleapis.com/auth/documents.readonly"]
#         creds = ServiceAccountCredentials.from_json_keyfile_name(
#             json_key_file, scope
#         )
#         service = build("docs", "v1", credentials=creds)
#         logger.info("✅ Connected to Google Doc successfully")
#         return service
#     except Exception as e:
#         logger.error(f"❌ Failed to connect to Google Doc: {e}")
#         raise

# # -------------------------
# # Google Sheets Handler
# # -------------------------
# class GoogleSheetsHandler:
#     def __init__(self, json_key_file, sheet_url):
#         self.json_key_file = json_key_file
#         self.sheet_url = sheet_url
#         self.sheet = _connect_to_google_sheet(json_key_file, sheet_url)

#     def _connect(self):
#         try:
#             scope = [
#                 "https://spreadsheets.google.com/feeds",
#                 "https://www.googleapis.com/auth/drive"
#             ]
#             creds = ServiceAccountCredentials.from_json_keyfile_name(
#                 self.json_key_file, scope
#             )
#             client = gspread.authorize(creds)
#             sheet = client.open_by_url(self.sheet_url).sheet1
#             logger.info("✅ Connected to Google Sheet successfully")
#             return sheet
#         except Exception as e:
#             logger.error(f"❌ Failed to connect to Google Sheet: {e}")
#             raise

#     def read(self):
#         try:
#             data = self.sheet.get_all_records()
#             logger.info("📊 Reading Google Sheet data")
#             for row in data:
#                 logger.info(row)
#             return data
#         except Exception as e:
#             logger.error(f"❌ Failed to read sheet: {e}")
#             raise

#     def write(self, row_data):
#         try:
#             self.sheet.append_row(row_data)
#             logger.info(f"✅ Row added: {row_data}")
#         except Exception as e:
#             logger.error(f"❌ Failed to write row: {e}")
#             raise

#     def delete(self, row_number):
#         try:
#             self.sheet.delete_rows(row_number)
#             logger.info(f"🗑️ Row {row_number} deleted successfully")
#         except Exception as e:
#             logger.error(f"❌ Failed to delete row {row_number}: {e}")
#             raise

# # -------------------------
# # Google Docs Handler
# # -------------------------
# class GoogleDocsHandler:
#     def __init__(self, json_key_file, doc_id):
#         self.json_key_file = json_key_file
#         self.doc_id = doc_id
#         self.service = _connect_to_google_doc(json_key_file, doc_id)

#     def read(self):
#         try:
#             document = self.service.documents().get(documentId=self.doc_id).execute()
#             title = document.get("title")
#             logger.info(f"📄 Google Doc Title: {title}")

#             logger.info("📄 Reading Google Doc content...")
#             text_content = []
#             for content in document.get("body", {}).get("content", []):
#                 if "paragraph" in content:
#                     elements = content["paragraph"]["elements"]
#                     for elem in elements:
#                         if "textRun" in elem:
#                             text_content.append(elem["textRun"]["content"])

#             full_text = "".join(text_content)
#             logger.info("✅ Document read successfully")
#             return full_text
#         except Exception as e:
#             logger.error(f"❌ Failed to read Google Doc: {e}")
#             raise

# # -------------------------
# # Main
# # -------------------------
# # def main():
# #     json_key_file = input("Enter path to your JSON key file: ").strip()
# #     sheet_url = input("Enter Google Sheet URL: ").strip()

# #     # Google Sheets operations
# #     sheets = GoogleSheetsHandler(json_key_file, sheet_url)
# #     # sheets.read()
# #     # sheets.write(["Charlie", "28", "Berlin"])
# #     # sheets.delete(2)
# #     sheets.read()

# #     # Google Docs operations
# #     doc_id = input("Enter Google Doc ID (from its URL): ").strip()
# #     docs = GoogleDocsHandler(json_key_file, doc_id)
# #     content = docs.read()
# #     print("\n--- Google Doc Content ---\n")
# #     print(content)

# # if __name__ == "__main__":
# #     main()




import os
import glob
import requests
from bs4 import BeautifulSoup
from docx import Document
import pandas as pd
import io
import re
import logging
import PyPDF2
import json

# LangChain / FAISS
from langchain_community.vectorstores import FAISS
from langchain_huggingface import HuggingFaceEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document as LCDocument

# HuggingFace tokenizer
from transformers import AutoTokenizer

logger = logging.getLogger(__name__)
logging.basicConfig(level=logging.INFO, format="%(message)s")

FAISS_INDEX_PATH = "faiss_index"


class VectorStore:
    def __init__(self, model_name: str = "sentence-transformers/all-MiniLM-L6-v2"):
        self.embeddings = HuggingFaceEmbeddings(model_name=model_name)
        self.vectors = None
        self.model_name = model_name
        self.tokenizer = AutoTokenizer.from_pretrained(model_name)

    # -------- HTML --------
    def extract_html(self, url: str) -> str:
        response = requests.get(url)
        soup = BeautifulSoup(response.text, "html.parser")
        return soup.get_text(separator="\n", strip=True)

    # -------- PDF --------
    def extract_pdf(self, url: str) -> str:
        response = requests.get(url)
        with io.BytesIO(response.content) as f:
            reader = PyPDF2.PdfReader(f)
            text = "\n".join(
                page.extract_text() for page in reader.pages if page.extract_text()
            )
        return text

    # -------- DOCX --------
    def extract_docx(self, url: str) -> str:
        response = requests.get(url)
        with io.BytesIO(response.content) as f:
            doc = Document(f)
            text = "\n".join(p.text for p in doc.paragraphs)
        return text

    # -------- XLSX --------
    def extract_xlsx(self, url: str) -> str:
        df = pd.read_excel(url)
        return df.to_string()

    # -------- Google Doc (Public) --------
    def extract_google_doc(self, doc_url: str) -> str:
        file_id = re.findall(r"/d/([a-zA-Z0-9-_]+)", doc_url)[0]
        export_url = f"https://docs.google.com/document/d/{file_id}/export?format=txt"
        response = requests.get(export_url)
        return response.text

    # -------- Google Sheet (Public) --------
    def extract_google_sheet(self, sheet_url: str) -> str:
        sheet_id = re.findall(r"/d/([a-zA-Z0-9-_]+)", sheet_url)[0]
        export_url = f"https://docs.google.com/spreadsheets/d/{sheet_id}/export?format=csv"
        response = requests.get(export_url)
        df = pd.read_csv(io.StringIO(response.text))
        return df.to_string()

    # -------- Dispatcher --------
    def extract(self, url: str) -> str:
        if "docs.google.com/document" in url:
            return self.extract_google_doc(url)
        elif "docs.google.com/spreadsheets" in url:
            return self.extract_google_sheet(url)
        elif url.endswith(".pdf"):
            return self.extract_pdf(url)
        elif url.endswith(".docx"):
            return self.extract_docx(url)
        elif url.endswith(".xlsx"):
            return self.extract_xlsx(url)
        else:
            return self.extract_html(url)

    # -------- Vector Store Management --------
    def ensure_index(self, url: str) -> dict:
        if self.vectors:
            logger.info("FAISS index already loaded.")
            return self._make_metadata(url, "FAISS index already loaded.", 0, 0)

        if os.path.exists(FAISS_INDEX_PATH) and glob.glob(f"{FAISS_INDEX_PATH}/*"):
            self.vectors = FAISS.load_local(
                FAISS_INDEX_PATH, self.embeddings, allow_dangerous_deserialization=True
            )
            logger.info("Loaded existing FAISS index.")
            return self._make_metadata(url, "Loaded existing FAISS index.", 0, 0)

        # Extract text directly (no file save)
        text = self.extract(url)
        logger.info("Text extracted successfully.")

        # Count total tokens
        total_tokens = len(self.tokenizer.encode(text))

        # Split text into chunks
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
        chunks = text_splitter.split_text(text)
        documents = [LCDocument(page_content=chunk) for chunk in chunks]

        # Count tokens after chunking
        embedded_tokens = sum(len(self.tokenizer.encode(chunk)) for chunk in chunks)

        # Create embeddings and FAISS index
        self.vectors = FAISS.from_documents(documents, self.embeddings)
        os.makedirs(FAISS_INDEX_PATH, exist_ok=True)
        self.vectors.save_local(FAISS_INDEX_PATH)

        logger.info("Vector embeddings created and saved.")
        logger.info(f"Total tokens in document: {total_tokens}")
        logger.info(f"Total tokens embedded: {embedded_tokens}")

        return self._make_metadata(url, "Vector embeddings created and saved.", total_tokens, embedded_tokens)

    def _make_metadata(self, url: str, status: str, total_tokens: int, embedded_tokens: int) -> dict:
        metadata = {
            "status": status,
            "source_url": url,
            "embedding_model": self.model_name,
            "vectorstore_path": FAISS_INDEX_PATH,
            "total_tokens": total_tokens,
            "embedded_tokens": embedded_tokens,
        }
        logger.info(json.dumps(metadata, indent=2))
        return metadata


# --------- TEST ---------
if __name__ == "__main__":
    store = VectorStore()
    test_url = "https://proceedings.neurips.cc/paper_files/paper/2017/file/3f5ee243547dee91fbd053c1c4a845aa-Paper.pdf"
    result = store.ensure_index(test_url)
    print(result)
